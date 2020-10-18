# Some random comment
import random, datetime, pytz, pycountry

from docx import Document

def formatDocument():
    document = Document()
    document.add_heading('Time Based Questions', 0)
    p = document.add_paragraph()
    p.add_run('This document contains time based questions.').bold = True

    return document

def getCountries():
    countryOrigin = countries[random.randrange(0,len(countries))]
    countryDest = countries[random.randrange(0,len(countries))]
    while countryOrigin == countryDest:
        countryOrigin = countries[random.randrange(0,len(countries))]
        countryDest = countries[random.randrange(0,len(countries))]

    return countryOrigin, countryDest

def getCountriesAndTimezones():
    countryOrigin, countryDest = getCountries()
    while True:
        try:
            tzOrigin = pytz.country_timezones[countryOrigin.alpha_2][0]
            tzDestination = pytz.country_timezones[countryDest.alpha_2][0]
            break
        except:
            countryOrigin, countryDest = getCountries()
            pass

    return countryOrigin, countryDest, tzOrigin, tzDestination

def getHour():
    return random.randrange(0,24)

def getMinute():
    return random.randrange(0,60)

def getFormatedDateTime(datetime):
    return datetime.strftime("%d/%m/%Y, %H:%M")

### Main Method Starts Here

numberOfQuestions = 100
countries = list(pycountry.countries)
document = formatDocument()

for i in range(numberOfQuestions):
    countryOrigin, countryDest, tzOrigin, tzDestination = getCountriesAndTimezones()

    hourOrigin = getHour()
    minuteOrigin = getMinute()
    utcOrigin = datetime.datetime(2020,3,4,hourOrigin, minuteOrigin, 0, 0, tzinfo = pytz.utc)

    dateTimeOrigin = utcOrigin.astimezone(pytz.timezone(tzOrigin))
    formattedOrigin = getFormatedDateTime(dateTimeOrigin)

    dateTimeDestination = utcOrigin.astimezone(pytz.timezone(tzDestination))
    formattedDestination = getFormatedDateTime(dateTimeDestination)

    randomHourOrigin = getHour()
    minuteOrigin = getMinute()
    randomUTCOrigin = datetime.datetime(2020,3,4,randomHourOrigin, minuteOrigin, 0, 0, tzinfo= pytz.utc)
    randomOriginDateTime = randomUTCOrigin.astimezone(pytz.timezone(tzOrigin))
    randomFormattedOrigin = getFormatedDateTime(randomOriginDateTime)

    durationHours = getHour()
    durationMinutes = getMinute()

    sectionOne = f'Question {i+1}:'
    sectionTwo = f'When the date and ime in {countryOrigin.name} is {formattedOrigin} respectively, ' \
        + f'the date and time in {countryDest.name} is {formattedDestination} respectively.'
    sectionThree = f' A plane leaves {countryOrigin.name} at {randomFormattedOrigin} local date and time and takes {durationHours}' \
        + f' hours and {durationMinutes} minutes to fly to {countryDest.name}.'
    sectionFour = f'What is the date and time in {countryDest.name} when it arrives?'
    answer = getFormatedDateTime((randomUTCOrigin + datetime.timedelta(0,0,0,0, durationMinutes, durationHours,0)).astimezone(pytz.timezone(tzDestination)))
    sectionFive = f'Answer: {answer}'

    document.add_paragraph(sectionOne)
    p = document.add_paragraph(sectionTwo)
    p.add_run(sectionThree)
    document.add_paragraph(sectionFour)
    document.add_paragraph(sectionFive)

document.save('test.docx')